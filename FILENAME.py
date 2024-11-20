import os
import io
import collections
import bson
from datetime import datetime, timezone

from bson.codec_options import CodecOptions
from pymongo.mongo_client import MongoClient
from pymongo.server_api import ServerApi
from googleapiclient.http import MediaIoBaseDownload
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from docx import Document

SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
folder_id = "14qABNrUF4fZ0aXMowjs6Bwkt-I9FmaGr"
dir = []
data = []

# performes all drive functions
def quickstart():

    # connects to drive
    creds = None
    if os.path.exists("token.json"):
        creds = Credentials.from_authorized_user_file("token.json", SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file("credentials.json", SCOPES)
            creds = flow.run_local_server(port=0)
        with open("token.json", "w") as token:
            token.write(creds.to_json())

    # accessing and downloading files from Feasibility Studies folder
    try:
        service = build("drive", "v3", credentials=creds)

        query = f"'{folder_id}' in parents and mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'"
        response = service.files().list(q=query, spaces="drive").execute()
        files = response.get("files", [])

        if not files:
            print("No .docx files found in the 'Feasibility' folder.")
            return []

        os.makedirs("downloads", exist_ok=True)

        for file in files:
            request = service.files().get_media(fileId=file["id"])
            file_path = os.path.join("downloads", f"{file['name']}.docx")
            with io.FileIO(file_path, "wb") as f:
                downloader = MediaIoBaseDownload(f, request)
                done = False
                while not done:
                    status, done = downloader.next_chunk()
                    print(
                        f"Downloading {file['name']}: {int(status.progress() * 100)}% complete."
                    )
            dir.append(file_path)

    except HttpError as error:
        print(f"An error occurred: {error}")


# connects and uploads files to mongodb
def upload_to_mongodb(uri):

    client = MongoClient(uri, server_api=ServerApi("1"))

    try:
        client.admin.command("ping")
        print("Pinged your deployment. You successfully connected to MongoDB!")
    except Exception as e:
        print(e)

    db = client["test_bibi"]
    collection = db["feasibility_studies"]

    for document in data:
        collection.insert_one(document)
        print("New document inserted!")

# parsing the docx files into bson
def get_files():

    for file in dir:
        document = Document(file)

        partner_organization_name = document.tables[0].rows[1].cells[1].text

        project_goal = document.tables[0].rows[2].cells[1].text

        # missing attributes Project Summary
        project_summary = "TODO"
        # missing attributes Project Name
        project_name = "TODO"
        # missing attributes Project ID
        project_ID = "TODO"

        project_coordinator = document.tables[0].rows[6].cells[1].text

        # missing attribute Project Contributor (?)
        project_contributor = document.tables[0].rows[5].cells[1].text.split(", ")
        # missing attribute Partner Organization Type
        partner_organization_type = "TODO"

        competence_group_involved = document.tables[0].rows[3].cells[1].text.split(", ")

        start_date = document.tables[0].rows[7].cells[1].text

        end_date = document.tables[0].rows[7].cells[3].text

        # missing attribute Operating Geo-region
        operating_geo_region = "TODO"
        # missing attribute Keywords
        keywords = ["TODO"]
        # missing attribute MEL-metrics
        mel_metrics = ["TODO"]

        data.append(
            {
                "Project name": project_name,
                "Project EWB-SWE ID:": project_ID,
                "Project goal": project_goal,
                "Project summary": project_summary,
                "Project coordinator": project_coordinator,
                "Project contributor": project_contributor,
                "Partner organization name": partner_organization_name,
                "Partner organization type": partner_organization_type,
                "Competence group involved": competence_group_involved,
                "Start date": start_date,
                "Complete date": end_date,
                "Operation geo-region": operating_geo_region,
                "Keywords": keywords,
                "MEL metrics": mel_metrics,
                "DocCreationTimeStamp": datetime.now(timezone.utc)
            }
        )

        print("New document added to the list!")


def main():

    quickstart()

    # asking user their credentials for mongodb
    print("Input mondodb credentials")
    username = input("Username:")
    password = input("Password:")
    uri = (
        "mongodb+srv://"
        + username
        + ":"
        + password
        + "@infoedge-cluster0.wrfbo.mongodb.net/"
    )

    get_files()
    upload_to_mongodb(uri)

if __name__ == "__main__":
    main()
